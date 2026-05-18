from __future__ import annotations

import asyncio
import logging
import os
import shutil
import tempfile
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.exceptions import RequestValidationError
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from paper_refiner.apply import replace_paragraph_plain_text
from paper_refiner.config_loader import load_filter_config, load_prompts
from paper_refiner.llm import (
    _llm_timeout,
    describe_llm_http_error,
    max_llm_concurrency,
    rewrite_text_async,
)
from paper_refiner.paths import project_root
from paper_refiner.prompt_template import inject_paragraph_text
from paper_refiner.scan import scan_document
from paper_refiner.web import sessions


@asynccontextmanager
async def _lifespan(_app: FastAPI):
    """加载项目根 ``.env``；``override=False`` 表示不覆盖已在命令行 / Shell 中设置的变量。"""
    env_file = project_root() / ".env"
    load_dotenv(env_file, override=False)
    load_dotenv(override=False)
    yield


HERE = Path(__file__).resolve().parent
MAX_UPLOAD_BYTES = 25 * 1024 * 1024
# 单次「生成改写」允许勾选的段落数上限（每段仍单独请求 LLM，串行处理）
MAX_REFINE_SELECTION = int(os.environ.get("PAPER_REFINER_MAX_BATCH", "1000"))

_log = logging.getLogger(__name__)

app = FastAPI(title="paper-refiner tool", version="0.2.0", lifespan=_lifespan)


@app.middleware("http")
async def _catch_unhandled(request: Request, call_next):
    try:
        return await call_next(request)
    except (HTTPException, RequestValidationError):
        raise
    except Exception as exc:  # noqa: BLE001 — 返回 JSON 便于前端展示，避免裸 Internal Server Error
        _log.exception("Unhandled error %s %s", request.method, request.url.path)
        msg = f"{type(exc).__name__}: {exc}"
        return JSONResponse(status_code=500, content={"detail": msg[:2500]})


app.mount("/static", StaticFiles(directory=str(HERE / "static")), name="static")


def _prompts_path() -> Path:
    return project_root() / "prompts.yaml"


def _config_path() -> Path:
    root = project_root()
    custom = root / "config.yaml"
    if custom.is_file():
        return custom
    return root / "config.example.yaml"


@app.get("/", response_class=HTMLResponse)
async def index() -> str:
    html = (HERE / "templates" / "index.html").read_text(encoding="utf-8")
    return html


@app.get("/api/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/api/config")
async def api_config() -> dict[str, Any]:
    """Non-secret defaults for the UI."""
    return {
        "max_refine_selection": MAX_REFINE_SELECTION,
        "max_llm_concurrent": max_llm_concurrency(),
        "default_model": os.environ.get("OPENAI_MODEL", "deepseek-v4-flash"),
        "api_base_hint": os.environ.get("OPENAI_API_BASE", "https://api.deepseek.com"),
    }


@app.get("/api/prompts")
async def api_prompts() -> dict[str, Any]:
    prompts = load_prompts(_prompts_path())
    return {"prompts": [{"id": p["id"], "name": p["name"]} for p in prompts]}


class RefineRequest(BaseModel):
    indices: list[int] = Field(..., min_length=1, max_length=MAX_REFINE_SELECTION)
    prompt_id: str
    temperature: float = Field(0.7, ge=0.0, le=2.0)
    model: str | None = None


class ExportEdit(BaseModel):
    index: int
    text: str


class ExportRequest(BaseModel):
    edits: list[ExportEdit] = Field(..., min_length=1)


class SelectionRequest(BaseModel):
    indices: list[int] = Field(default_factory=list, max_length=MAX_REFINE_SELECTION)
    prompt_id: str | None = None
    temperature: float | None = Field(None, ge=0.0, le=2.0)
    model: str | None = None


@app.post("/api/upload")
async def upload(file: UploadFile = File(...)) -> dict[str, Any]:
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "请上传 .docx 文件")
    raw = await file.read()
    if len(raw) > MAX_UPLOAD_BYTES:
        raise HTTPException(400, f"文件过大（上限 {MAX_UPLOAD_BYTES // (1024 * 1024)} MB）")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(raw)
        tmp_path = tmp.name
    try:
        fcfg = load_filter_config(_config_path())
        infos = scan_document(tmp_path, fcfg)
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    sid = sessions.create_session(raw, infos, original_filename=file.filename)
    meta = sessions.load_meta(sid)
    return {
        "session_id": sid,
        "paragraphs": meta["paragraphs"],
        "eligible_indices": meta["eligible_indices"],
        "selected_indices": meta.get("selected_indices") or [],
        "refine_options": meta.get("refine_options") or {},
        "original_filename": meta.get("original_filename", file.filename),
    }


@app.get("/api/sessions/recent")
async def recent_sessions() -> dict[str, Any]:
    return {"sessions": sessions.list_recent_sessions()}


@app.get("/api/session/{session_id}")
async def get_session(session_id: str) -> dict[str, Any]:
    try:
        return sessions.session_payload(session_id)
    except (ValueError, FileNotFoundError):
        raise HTTPException(404, "会话不存在或已过期")


@app.put("/api/session/{session_id}/selection")
async def save_selection(session_id: str, body: SelectionRequest) -> dict[str, Any]:
    try:
        meta = sessions.load_meta(session_id)
    except (ValueError, FileNotFoundError):
        raise HTTPException(404, "会话不存在或已过期")

    eligible = set(meta["eligible_indices"])
    ordered = sorted(set(body.indices))
    if len(ordered) != len(body.indices):
        raise HTTPException(400, "段落索引不能重复")
    for idx in ordered:
        if idx not in eligible:
            raise HTTPException(400, f"段落 {idx} 不可改写")

    refine_options: dict[str, Any] = {}
    if body.prompt_id is not None:
        refine_options["prompt_id"] = body.prompt_id
    if body.temperature is not None:
        refine_options["temperature"] = body.temperature
    if body.model is not None:
        refine_options["model"] = body.model.strip() or None

    sessions.save_selection(
        session_id,
        selected_indices=ordered,
        refine_options=refine_options or None,
    )
    return {
        "ok": True,
        "selected_indices": ordered,
        "refine_options": sessions.load_meta(session_id).get("refine_options") or {},
    }


@app.post("/api/session/{session_id}/refine")
async def refine(session_id: str, body: RefineRequest) -> dict[str, Any]:
    try:
        meta = sessions.load_meta(session_id)
    except (ValueError, FileNotFoundError):
        raise HTTPException(404, "会话不存在或已过期")

    eligible = set(meta["eligible_indices"])
    by_index = {p["index"]: p for p in meta["paragraphs"]}

    ordered = sorted(set(body.indices))
    if len(ordered) != len(body.indices):
        raise HTTPException(400, "段落索引不能重复")

    for idx in ordered:
        if idx not in eligible:
            raise HTTPException(400, f"段落 {idx} 不可改写（含公式/图/样式限制等）")

    prompts = load_prompts(_prompts_path())
    by_id = {p["id"]: p for p in prompts}
    if body.prompt_id not in by_id:
        raise HTTPException(400, f"未知 prompt_id: {body.prompt_id}")

    orig = sessions.original_path(session_id)
    doc = Document(str(orig))

    sem = asyncio.Semaphore(max_llm_concurrency())
    conn_limit = max_llm_concurrency() + 10
    client_limits = httpx.Limits(
        max_keepalive_connections=conn_limit, max_connections=conn_limit
    )

    async def rewrite_one(idx: int, http: httpx.AsyncClient) -> dict[str, Any]:
        text = doc.paragraphs[idx].text or ""
        template = by_id[body.prompt_id]["template"]
        try:
            msg = inject_paragraph_text(template, text)
        except ValueError as e:
            raise HTTPException(500, str(e)) from e
        try:
            async with sem:
                after = await rewrite_text_async(
                    msg,
                    model=body.model,
                    temperature=body.temperature,
                    client=http,
                )
        except RuntimeError as e:
            msg = str(e)
            extra = ""
            if "未设置" in msg or "Missing environment variable" in msg:
                extra = " 请确认项目根目录存在 `.env`，且已重启 `paper-refiner-web`。"
            raise HTTPException(503, msg + extra) from e
        except httpx.HTTPStatusError as e:
            sc = e.response.status_code if e.response else 0
            provider_msg = describe_llm_http_error(e.response)
            if sc == 401:
                raise HTTPException(
                    status_code=502,
                    detail=(
                        "大模型 API 认证失败（HTTP 401）：密钥无效或未生效。\n"
                        "请逐项检查：\n"
                        "1）在 https://platform.deepseek.com/api_keys 创建并复制 **完整** API Key；"
                        "使用命令行启动：`paper-refiner-web --api-key sk-... --api-base https://api.deepseek.com`，"
                        "或在项目根目录 `.env` 中：`OPENAI_API_KEY=sk-...`（不要加引号）。\n"
                        "   也可使用别名：DEEPSEEK_API_KEY=sk-xxxxxxxx\n"
                        "2）`OPENAI_API_BASE` 须与密钥同一服务商，DeepSeek 官方为：\n"
                        "   OPENAI_API_BASE=https://api.deepseek.com\n"
                        "   （程序会自动补全为 …/v1 再请求 /chat/completions；也可用命令行 --api-base）\n"
                        "3）修改 `.env` 后 **必须重启** `paper-refiner-web` 进程，否则仍用旧环境变量。\n"
                        "4）若从别处复制，注意不要混入中文全角字符或漏复制首尾字符。\n"
                        f"服务商摘要：{provider_msg or '(无正文)'}"
                    ),
                ) from e
            raise HTTPException(
                502,
                f"大模型 API 返回 HTTP {sc}: {provider_msg or e.response.text[:800] if e.response else str(e)}",
            ) from e
        except httpx.RequestError as e:
            retries = os.environ.get("OPENAI_RETRY_COUNT", "3")
            raise HTTPException(
                502,
                f"无法连接大模型 API（已自动重试最多 {retries} 次）: {e}",
            ) from e
        meta_p = by_id[body.prompt_id]
        return {
            "index": idx,
            "before": text,
            "after": after,
            "style": by_index[idx].get("style", ""),
            "prompt_id": body.prompt_id,
            "prompt_name": meta_p.get("name", body.prompt_id),
        }

    async def run_parallel() -> list[dict[str, Any]]:
        async with httpx.AsyncClient(
            timeout=_llm_timeout(), limits=client_limits
        ) as http_client:
            return await asyncio.gather(
                *[rewrite_one(idx, http_client) for idx in ordered]
            )

    try:
        results = list(await run_parallel())
    except HTTPException:
        raise
    sessions.save_last_results(session_id, results)
    return {"results": results}


@app.post("/api/session/{session_id}/export")
async def export_docx(session_id: str, body: ExportRequest) -> FileResponse:
    try:
        meta = sessions.load_meta(session_id)
    except (ValueError, FileNotFoundError):
        raise HTTPException(404, "会话不存在或已过期")

    eligible = set(meta["eligible_indices"])
    n_para = len(meta["paragraphs"])
    indices_seen = set()
    for e in body.edits:
        if e.index in indices_seen:
            raise HTTPException(400, f"段落 {e.index} 在导出列表中重复")
        indices_seen.add(e.index)
        if e.index not in eligible:
            raise HTTPException(400, f"段落 {e.index} 不允许写入")
        if e.index < 0 or e.index >= n_para:
            raise HTTPException(400, f"段落索引越界: {e.index}")

    orig = sessions.original_path(session_id)
    d = sessions.session_dir(session_id)
    out = d / "export_latest.docx"
    shutil.copyfile(orig, out)

    doc = Document(str(out))
    for e in body.edits:
        replace_paragraph_plain_text(doc.paragraphs[e.index], e.text)
    doc.save(str(out))

    return FileResponse(
        path=str(out),
        filename="refined.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.delete("/api/session/{session_id}")
async def discard_session(session_id: str) -> dict[str, str]:
    try:
        sessions.delete_session(session_id)
    except ValueError:
        raise HTTPException(400, "invalid session")
    return {"ok": "true"}
