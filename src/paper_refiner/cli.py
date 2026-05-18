from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Optional

import httpx
import typer
from dotenv import load_dotenv

from paper_refiner.apply import replace_paragraph_plain_text
from paper_refiner.config_loader import load_filter_config, load_prompts
from paper_refiner.llm import _llm_timeout, max_llm_concurrency, rewrite_text_async
from paper_refiner.paths import project_root
from paper_refiner.prompt_template import inject_paragraph_text
from paper_refiner.scan import eligible_indices, scan_document

app = typer.Typer(no_args_is_help=True, add_completion=False)


def _bootstrap_env() -> None:
    """与 Web 一致：优先加载项目根目录 `.env`，避免 cwd 不在仓库目录时读不到密钥。"""
    load_dotenv(project_root() / ".env", override=True)
    load_dotenv(override=False)


@app.callback()
def _main_callback() -> None:
    _bootstrap_env()


def _default_paths() -> tuple[Path, Path]:
    root = project_root()
    return root / "prompts.yaml", root / "config.example.yaml"


@app.command("list-doc")
def list_doc(
    docx: Path = typer.Argument(..., exists=True, readable=True, help="Path to .docx"),
    prompts_file: Optional[Path] = typer.Option(None, help="prompts.yaml"),
    config: Optional[Path] = typer.Option(None, help="filter config yaml"),
    eligible_only: bool = typer.Option(False, help="Only show paragraphs that can be refined"),
) -> None:
    """List all paragraphs with skip reasons (formulas/images/tables/styles)."""
    pfile, cdefault = _default_paths()
    prompts_path = prompts_file or pfile
    cfg_path = config or cdefault
    load_prompts(prompts_path)  # validate early
    fcfg = load_filter_config(cfg_path)
    infos = scan_document(str(docx), fcfg)
    typer.echo(f"Document: {docx}")
    typer.echo(f"Filter config: {cfg_path}")
    for info in infos:
        if eligible_only and info.skip_reason:
            continue
        flag = "" if info.skip_reason else "OK"
        reason = info.skip_reason or flag
        typer.echo(
            f"[{info.index:4d}] {reason:22s} | style={info.style!r} | {info.preview}"
        )
    typer.echo("")
    typer.echo(f"Eligible paragraph indices: {eligible_indices(infos)}")


@app.command("prompts")
def list_prompts(
    prompts_file: Optional[Path] = typer.Option(None, help="prompts.yaml"),
) -> None:
    """Show prompt ids loaded from prompts.yaml."""
    pfile, _ = _default_paths()
    prompts_path = prompts_file or pfile
    prompts = load_prompts(prompts_path)
    for p in prompts:
        typer.echo(f"- {p['id']}: {p['name']}")


@app.command("rewrite")
def rewrite_cmd(
    docx: Path = typer.Argument(..., exists=True, readable=True),
    indices: str = typer.Option(
        ...,
        help="Comma-separated paragraph indices, e.g. '12' or '12,13' (keep small)",
    ),
    prompt_id: str = typer.Option(..., help="Prompt id from prompts.yaml"),
    prompts_file: Optional[Path] = typer.Option(None),
    config: Optional[Path] = typer.Option(None),
    output: Optional[Path] = typer.Option(
        None,
        help="Write updated docx here (required unless --dry-run)",
    ),
    dry_run: bool = typer.Option(False, help="Print model output only; do not modify file"),
    max_paragraphs: int = typer.Option(
        1000,
        help="单次命令最多处理的段落数（每段单独请求 LLM，按编号升序串行）",
    ),
    temperature: float = typer.Option(0.7, min=0.0, max=2.0),
    model: Optional[str] = typer.Option(None, help="Override OPENAI_MODEL"),
) -> None:
    """Rewrite selected plain-text body paragraphs using the chosen prompt."""
    pfile, cdefault = _default_paths()
    prompts_path = prompts_file or pfile
    cfg_path = config or cdefault
    prompts = load_prompts(prompts_path)
    by_id = {p["id"]: p for p in prompts}
    if prompt_id not in by_id:
        raise typer.BadParameter(f"Unknown prompt_id {prompt_id!r}; try: paper-refiner prompts")

    idxs = sorted({int(x.strip()) for x in indices.split(",") if x.strip()})
    if len(idxs) > max_paragraphs:
        raise typer.BadParameter(f"Too many indices ({len(idxs)}); max is {max_paragraphs}")

    fcfg = load_filter_config(cfg_path)
    infos = scan_document(str(docx), fcfg)
    by_index = {i.index: i for i in infos}

    for idx in idxs:
        if idx not in by_index:
            raise typer.BadParameter(f"Index {idx} out of range")
        info = by_index[idx]
        if info.skip_reason:
            raise typer.BadParameter(
                f"Paragraph {idx} is not eligible ({info.skip_reason}). "
                "Choose another index or adjust filters."
            )

    from docx import Document

    doc = Document(str(docx))

    async def run_parallel() -> list[tuple[int, str]]:
        conc = max_llm_concurrency()
        sem = asyncio.Semaphore(conc)
        conn_lim = conc + 10
        limits = httpx.Limits(
            max_keepalive_connections=conn_lim, max_connections=conn_lim
        )
        template = by_id[prompt_id]["template"]

        async def one(client: httpx.AsyncClient, idx: int) -> tuple[int, str]:
            text = doc.paragraphs[idx].text or ""
            msg = inject_paragraph_text(template, text)
            async with sem:
                rewritten = await rewrite_text_async(
                    msg, model=model, temperature=temperature, client=client
                )
            return (idx, rewritten)

        async def run() -> list[tuple[int, str]]:
            async with httpx.AsyncClient(timeout=_llm_timeout(), limits=limits) as client:
                return list(
                    await asyncio.gather(*(one(client, idx) for idx in idxs))
                )

        return await run()

    pairs = asyncio.run(run_parallel())

    for idx, new_text in pairs:
        typer.echo("-----")
        typer.echo(f"Paragraph {idx} — prompt: {prompt_id}")
        typer.echo(f"Paragraph {idx} — before:\n{by_index[idx].full_text}\n")
        typer.echo(f"Paragraph {idx} — after:\n{new_text}\n")

    if dry_run:
        typer.echo("(dry-run: file not modified)")
        return

    if output is None:
        raise typer.BadParameter("Pass --output path.docx or use --dry-run")

    for idx, new_text in pairs:
        replace_paragraph_plain_text(doc.paragraphs[idx], new_text)

    doc.save(str(output))
    typer.echo(f"Saved: {output}")


def main() -> None:
    app()


if __name__ == "__main__":
    main()
